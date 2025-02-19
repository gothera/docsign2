from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE 
from typing import Optional, Union, List, Tuple
from fuzzywuzzy import fuzz
import sys
import os

rootPath = os.getenv('PATH_TO_ROOT')

threshold = 98

class DocEditor:

    def __init__(self, document: Union[Document, str] = None, outputFileName: Optional[str] = '', debug: Optional[bool] = False):
        if document is None:
            raise ValueError("document cannot be None")
        if isinstance(document, str):
            if not outputFileName:
                outputFileName = document[:-5] + '_output.docx'
            document = Document(document)
        
        self.debug = debug
        self.outputFileName = outputFileName
        self.document: Document = document

    def __barbaricSplit(self, no: int, text: str) -> List[str]:
        split = []
        for i in range(no):
            start, end = i * len(text) // no, (i + 1) * len(text) // no
            split.append(text[start:end])
        
        return split

    
    def replaceText(self, oldText: str, newText: str, save: bool = False) -> Document:
        if not oldText:
            if newText:
                raise ValueError('oldText can not be empty')
            return doc
        if save and not self.outputFileName:
            raise ValueError('The document can not be save if an outputFileName was not provided in the constructor')
        
        if self.debug:
            print(f'oldText: {oldText} newText: {newText}')
        
        paragraphsIndexes, offset = [], 0
        try:
            paragraphsIndexes, offset = self.find(oldText)
        except ValueError:
            if self.debug:
                print('Did not find exact match, trying fuzzyFind')
            paragraphsIndexes = self.findFuzzy(oldText)
       
        if len(paragraphsIndexes) == 0:
            raise RuntimeError(f'The list of paragraphs contianing the text is empty')

        doc = self.document
        if offset != 0:
            if self.debug:
                paragraphText = doc.paragraphs[paragraphsIndexes[0]].text
                print(f'old text start in paragraph {paragraphsIndexes[0]}: "{paragraphText}" at offset {offset} ("{paragraphText[offset:]}")')
            raise ValueError('oldText starts in the midle of a paragraph')

        # hopefully the most common case:
        if len(paragraphsIndexes) == 1:
            doc.paragraphs[paragraphsIndexes[0]].text = newText
            if save:
                doc.save(self.outputFileName)
            return doc
        
        paragraphsToChange = [doc.paragraphs[i] for i in paragraphsIndexes]
        firstParaStyle = paragraphsToChange[0].style
        allSameStyle = True
        for paragraph in paragraphsToChange[1:]:
            if paragraph.style.name != firstParaStyle.name:
                allSameStyle = False
                break
        
        if not allSameStyle:
            raise NotImplementedError('Not implemented methode to change multiple paragraphs with difrent styles')
        
        # TODO: implement a nicer split
        newtextChuncks  = self.__barbaricSplit(len(paragraphsToChange), newText)
        for newTextChunk, paragraph in zip(newtextChuncks, paragraphsToChange):
            paragraph.text = newTextChunk

        if save:
            doc.save(self.outputFileName)

        return doc

    def find(self, text: str) -> Tuple[List[int], int]:
        doc: Document = self.document
        partialSum = 0
        indexes = [partialSum]
        for p in doc.paragraphs:
            # in between every paragraph is a '\n' character as separator
            partialSum += len(p.text) + 1
            indexes.append(partialSum)
        
        if self.debug:
            print('List of starting indeces: ', indexes) 

        fullText = '\n'.join([p.text for p in doc.paragraphs])
        indexInText = fullText.find(text)

        if indexInText == -1:
            raise ValueError(f"The substring '{text}' was not found in the Document")            

        if self.debug:
            print(f'index in full text: {indexInText}, len text: {len(text)}')

        indexFirstParagraph, indexLastParagraph = -1, -1
        for i, index in enumerate(indexes):
            if index <= indexInText < indexes[i+1]:
                indexFirstParagraph = i
                break

        if indexFirstParagraph != -1:
            for i in range(indexFirstParagraph, len(indexes)):
                if indexes[i] < indexInText + len(text) <= indexes[i+1]:
                    indexLastParagraph = i
                    break

        if indexFirstParagraph == -1 or indexLastParagraph == -1:
            raise RuntimeError(f'The paragraphs contianing the text was not found (indexFirstParagraph={indexFirstParagraph}, indexLastParagraph={indexLastParagraph})')
        
        offset = indexInText - indexes[indexFirstParagraph]
        if self.debug:
            print(f'indexFirstParagraph={indexFirstParagraph}, indexLastParagraph={indexLastParagraph}, offset={offset}')

        return list(range(indexFirstParagraph, indexLastParagraph + 1)), offset
    
    def findFuzzy(self, text: str) -> List[int]:
        doc: Document = self.document
        partialSum = 0
        indexes = [partialSum]
        for p in doc.paragraphs:
            # in between every paragraph is a '\n' character as separator
            partialSum += len(p.text) + 1
            indexes.append(partialSum)
        
        if self.debug:
            print('List of starting indeces: ', indexes) 

        fullText = '\n'.join([p.text for p in doc.paragraphs])
        posibleMatches = []
        for index in indexes:
            score = fuzz.ratio(fullText[index: index + len(text)], text)
            if score >= threshold:
                posibleMatches.append((score, index))
        
        posibleMatches.sort(reverse=True)
        if not posibleMatches:
            raise ValueError(f"The substring '{text}' was not found in a fuzzy way in the Document with score above {threshold}")            

        if self.debug or len(posibleMatches) > 1:
            print(f'found {len(posibleMatches)} possible matches', file=sys.stderr if len(posibleMatches) > 1 else sys.stdout)
            for score, matchIndex in posibleMatches:
                print(f'Possible matche at index {matchIndex}: {fullText[matchIndex: matchIndex + len(text)]}')
        
        indexInText = posibleMatches[0][1]

        indexFirstParagraph, indexLastParagraph = -1, -1
        for i, index in enumerate(indexes):
            if index <= indexInText < indexes[i+1]:
                indexFirstParagraph = i
                break

        if indexFirstParagraph != -1:
            for i in range(indexFirstParagraph, len(indexes)):
                if indexes[i] < indexInText + len(text) <= indexes[i+1]:
                    indexLastParagraph = i
                    break

        if indexFirstParagraph == -1 or indexLastParagraph == -1:
            raise RuntimeError(f'The paragraphs contianing the text was not found (indexFirstParagraph={indexFirstParagraph}, indexLastParagraph={indexLastParagraph})')
        
        offset = indexInText - indexes[indexFirstParagraph]
        if self.debug:
            print(f'indexFirstParagraph={indexFirstParagraph}, indexLastParagraph={indexLastParagraph}, offset={offset}')

        return list(range(indexFirstParagraph, indexLastParagraph + 1))
    
    def removeText(self, text: str, save: Optional[bool] = False) -> Document:
        if not text:
            raise ValueError('Text can not be empty')
        if save and not self.outputFileName:
            raise ValueError('The document can not be save if an outputFileName was not provided in the constructor')

        paragraphsIndexes, offset = [], 0
        try:
            paragraphsIndexes, offset = self.find(text)
        except ValueError:
            if self.debug:
                print('Did not find exact match, trying fuzzyFind')
            paragraphsIndexes = self.findFuzzy(text)
       
        if len(paragraphsIndexes) == 0:
            raise RuntimeError(f'The list of paragraphs contianing the text is empty')
        
        if self.debug:
            print(paragraphsIndexes, len(doc.paragraphs), offset)

        doc = self.document
        index = 0
        for i in paragraphsIndexes:
            if i == paragraphsIndexes[0] and offset != 0:
                doc.paragraphs[i].text = doc.paragraphs[i].text[:offset]
                continue
            if i == paragraphsIndexes[-1]:
                if len(text) - index == len(doc.paragraphs[i].text):
                    # doc.paragraphs.remove(doc.paragraphs[i])
                    # doc.paragraphs.pop(i)
                    doc.paragraphs[i].text = ''
                else:
                    doc.paragraphs[i].text = doc.paragraphs[i].text[len(text) - index:]
                print(len(doc.paragraphs))
                continue
            # doc.paragraphs.remove(doc.paragraphs[i])
            doc.paragraphs[i].text = ''
            index += len(doc.paragraphs[i].text)

        if save:
            doc.save(self.outputFileName)
        
        return doc
    
    def addParagraph(self, textBefore: str, paragraphText: str, save: Optional[bool] = False) -> Document:
        if not textBefore:
            raise ValueError('TextBefore can not be empty')
        if not paragraphText:
            raise ValueError('paragraphText can not be empty')
        if save and not self.outputFileName:
            raise ValueError('The document can not be save if an outputFileName was not provided in the constructor')

        paragraphsIndexes = []
        try:
            paragraphsIndexes, _ = self.find(textBefore)
        except ValueError:
            if self.debug:
                print('Did not find exact match, trying fuzzyFind')
            paragraphsIndexes = self.findFuzzy(textBefore)
       
        if len(paragraphsIndexes) == 0:
            raise RuntimeError(f'The list of paragraphs contianing the text is empty')
        
        if self.debug:
            print(paragraphsIndexes, len(doc.paragraphs), offset)

        doc = self.document
        doc = Document()
        lastParagraph = doc.paragraphs[paragraphsIndexes[-1]]
        if (len(textBefore) >= len(lastParagraph.text) and fuzz.ratio(lastParagraph.text, textBefore[-len(lastParagraph.text):]) > threshold) or \
            (len(textBefore) < len(lastParagraph.text) and fuzz.ratio(textBefore, lastParagraph.text[-len(textBefore):]) > threshold):
            
            # add paragraphs after the last found paragraph
            text, style = paragraphText, doc.styles['Normal']
            nextText, nextStyle = '', None
            needsNewParagraph = True
            for i in range(paragraphsIndexes[-1] + 1, len(doc.paragraphs)):
                nextText, nextStyle = doc.paragraphs[i].text, doc.paragraphs[i].style
                doc.paragraphs[i].text, doc.paragraphs[i].style = text, style
                text, style = nextText, nextStyle
                if nextText.strip() == '':
                    # incorporate a new line in the previous paragraph
                    if '\n' in newText:
                        doc.paragraphs[i].text += '\n' 
                    needsNewParagraph = False
                    break
            if needsNewParagraph:
                doc.add_paragraph(text, style)

        else:
            # add text in the midle of the last paragraph found
            raise NotImplementedError('add text in the midle of the last paragraph found')


class TestDocEditor:
    def __init__(self, debug = False):
        self.filename = os.path.join(rootPath, 'data/testLongText.docx')
        self.filename1 = os.path.join(rootPath, 'data/startup.docx')
        self.document = Document(self.filename)
        self.docEditor = DocEditor(document=self.filename, debug=debug)
    
    def testReplaceTextSigleParagraphs(self):
        targetIndex = 4
        oldText = self.document.paragraphs[targetIndex].text
        newText = 'This text was replaced'
        editedDoc = self.docEditor.replaceText(oldText, newText, save=True)
        assert editedDoc.paragraphs[targetIndex].text == newText, f'Paragraph {targetIndex} from the doc shoud be "{newText}" but is "{editedDoc.paragraphs[targetIndex].text}"\n"'

    def testReplaceTextMultipleParagraphs(self):
        targetIndexes = [8, 9, 10]
        newText = 'This text was replaced\nThis text was replaced\nThis text was replaced!'
        oldText = '\n'.join([self.document.paragraphs[i].text for i in targetIndexes])
        editedDoc = self.docEditor.replaceText(oldText, newText, save=True)
         
        # use __barbaric split
        no = len(targetIndexes)
        for i, targetIndex in enumerate(targetIndexes):
            start, end = i * len(newText) // no, (i + 1) * len(newText) // no
            assert editedDoc.paragraphs[targetIndex].text == newText[start:end], f'Paragraph {targetIndex} from the doc shoud be "{newText[start:end]}" but is "{editedDoc.paragraphs[targetIndex].text}"\n"'

    # test replace with fuzzy

    def testFindText(self):
        targetIndexes = [8, 9, 10]
        text = '\n'.join([self.document.paragraphs[i].text for i in targetIndexes])
        listIndexes, offset = self.docEditor.find(text)
        assert offset == 0, f'offset shoud be 0 but is {offset}'
        assert listIndexes == targetIndexes, f'listIndexes shoud be {targetIndexes} but is {listIndexes}'
    
    def testFindFuzzyText(self):
        targetIndexes = [8, 9, 10]
        text = '\n'.join([self.document.paragraphs[i].text for i in targetIndexes])
        listIndexes = self.docEditor.findFuzzy(text)
        assert listIndexes == targetIndexes, f'listIndexes shoud be {targetIndexes} but is {listIndexes}'
    
    def testDelete(self):
        targetIndex = 4
        text = self.document.paragraphs[targetIndex].text
        doc = self.docEditor.removeText(text, save=True)
        for i, paragraph in enumerate(doc.paragraphs):
            assert text not in paragraph.text, f'a paragraph with remove text is still presesnt: {i} {paragraph.text}'

    # test deleteing from multimple paragraphs
    # test delete with fuzzy

    # test add simple find with 1 paragraphs as text before
    # test add simple find with 2 paragraphs as text before
    # test add fuzzy find with 2 paragraphs before


class printdocs:
    @staticmethod
    def print_text_with_info(document):
        for paragraph in document.paragraphs:
            print(f'{paragraph.text} (style: {paragraph.style}, alignment: {paragraph.alignment}')

    @staticmethod
    def print_text(document):
        for i, paragraph in enumerate(document.paragraphs):
            print(f'{i}. {paragraph.text}')

    @staticmethod
    def print_styles(document, filter=None):
        l = document.styles
        if filter:
            l = [s for s in l if s.type == filter]
        for style in l:
            print(f'Style: {style.name}  style type: {style.type}')

    @staticmethod
    def print_table(document):
        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    print(cell.text)

    @staticmethod            
    def print_header_footer(document):
        for section in document.sections:
            print(section.header.paragraphs[0].text if section.header.paragraphs else 'No header')
            print(section.footer.paragraphs[0].text if section.footer.paragraphs else 'No footer')


if __name__ == '__main__':
    tester = TestDocEditor(debug=True)
    # tester.testFindText()
    # tester.testReplaceTextSigleParagraphs()
    # tester.testReplaceTextMultipleParagraphs()
    # tester.testDelete()
    # tester.testFindFuzzyText()
    
    file = os.path.join(rootPath, 'data/startup.docx')
    original = "Counter-intuitively, most of the best startup ideas already have competitors, and founders incorrectly shy away from spaces with competitors.  It’s often a bigger reason to worry if you have zero competitors - that may mean that there is no need for this product (a SISP).  If your competitors are new or don’t have much marketshare, you can often just ignore them.\nBut if you are going up against an entrenched competitor (i.e., you want to beat Google at web search), you’re going to need a specific strategy to do that."
    newText = "  lkvbajks;;dbvqe;jdvba;osldjvbnas;oudvbqw;ijvklbnas;jdk;vl/asndbvhia;sjldhvnb ;adshkjlvkba;shkjlvba shkjbdlv"
    
    editor = DocEditor(file, debug=True)
    editor.replaceText(original, newText, save=True)

